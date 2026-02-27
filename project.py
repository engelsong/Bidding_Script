from docx import Document


class Project(object):
    """通过Word文档建立项目对象保存项目信息"""

    def __init__(self, document_name):
        self.name = None  # 项目名称
        self.code = None  # 招标编号
        self.date = None  # 开标日期
        self.destination = None  # 运抵目的地
        self.trans = None  # 运输方式
        self.trans_time = None  # 发运时间
        self.totalsum = 0  # 对外货值
        self.is_lowprice = False  # 是否为低价法
        self.sec_comlist = False  # 是否有供货清单二
        self.is_tech = False  # 是否有技术服务
        self.is_qa = False  # 是否有售后
        self.is_cc = False  # 是否来华培训
        self.techinfo = []  # 存放技术服务信息，格式为[人数，天数]
        self.training_days = 0  # 来华培训天数
        self.training_num = 0  # 来华培训人数
        self.qc = []  # 法检物资序号
        self.main_item = [] # 主要标的
        self.commodities = {}  # 存放物资信息字典
        self.commodities2 = {}  # 存放供货清单二物资
        self.demand_info = []
        document = Document(document_name)
        table_info = document.tables[0]  # 读取项目基本信息
        table_item1 = document.tables[1]  # 读取供货清单1
        table_demand = document.tables[2]   # 读取服务需求
        project_info = []
        for cell in table_info.column_cells(1):
            project_info.append(cell.text)

        table_demand_length = len(table_demand.rows)
        for index in range(table_demand_length):  # 从第1行开始读取表格
            temp = []
            row_now = table_demand.row_cells(index)
            for cell in row_now:
                temp.append(cell.text.strip())
            self.demand_info.append(temp)          

        table_item1_length = len(table_item1.rows)
        for index in range(1, table_item1_length):  # 从第2行开始读取表格
            temp = []
            row_now = table_item1.row_cells(index)
            length_row = len(row_now)
            for i in range(1, length_row):  # 将每行信息放入暂存数组
                temp.append(row_now[i].text.strip())
            temp.append(row_now[0].text.strip())  # 把物资编号放在最后一位
            self.commodities[index] = temp

        self.name, self.code, self.date, self.destination, self.trans, self.trans_time = project_info[0:6]
        self.totalsum = int(project_info[6])

        if project_info[7] in 'yY':
            self.is_lowprice = True
        if project_info[8] in 'yY':
            self.sec_comlist = True
            # table3 = document.tables[2]
            # self.commodities2 = {}  # 存放供货清单二物资
            # # 读取供货清单二
            # table3_length = len(table3.rows)
            # for index in range(1, table3_length):  # 从第2行开始读取表格
            #     temp = []
            #     row_now = table3.row_cells(index)
            #     length_row = len(row_now)
            #     for i in range(1, length_row - 1):  # 将每行信息放入暂存数组
            #         if i == 6:
            #             amount = ''
            #             the_unit = ''
            #             for d in row_now[i].text:
            #                 if d.isdigit():
            #                     amount += d
            #             the_unit = row_now[i].text.replace(amount, '')
            #             temp.append(amount)
            #             temp.append(the_unit)
            #         else:
            #             temp.append(row_now[i].text)
            #     price = ''
            #     for d in row_now[length_row - 1].text:
            #         if d.isdigit() or d == '.':
            #             price += d
            #     temp.append(float(price))  # 将金额转换为float
            #     temp.append(row_now[0].text)  # 把物资编号放在最后一位
            #     self.commodities2[index] = temp

        if project_info[9] in 'yY':
            self.is_tech = True
            self.techinfo += list(map(int, project_info[10:12]))
        if project_info[12] in 'yY':
            self.is_qa = True
        if project_info[13] in 'yY':
            self.is_cc = True
            self.training_days = int(project_info[15])  # 读取来华陪训天数
            self.training_num = int(project_info[14])  # 读取来华培训人数
        if project_info[-2] != '':
            if project_info[-1] not in 'Nn':
                self.qc += list(map(int, project_info[-1].split()))
                self.qc.sort()
        self.main_item = list(map(int, project_info[-1].split()))

    def show_info(self):
        print('项目名称:', self.name)
        print('项目代码:', self.code)
        print('开标日期:', self.date)
        print('目的地:', self.destination)
        print('运输方式:', self.trans)
        print('运输时间:', self.trans_time)
        print('对外货值：', self.totalsum)
        print('是否为低价法', '是' if self.is_lowprice is True else '否')
        print('是否有供货清单二', '是' if self.sec_comlist is True else '否')
        print('是否有技术服务:', '是' if self.is_tech is True else '否')
        print('是否有售后服务:', '是' if self.is_qa is True else '否')
        print('是否有来华培训', '是' if self.is_cc is True else '否')
        if self.is_tech:
            print('技术服务人数:', self.techinfo[0])
            print('技术服务天数:', self.techinfo[1])
        if self.is_cc:
            print('来华培训人数：', self.training_num)
            print('来华培训天数：', self.training_days)
        if len(self.qc) > 0:
            print('法检物资：', self.qc)
        print('主要标的有', self.main_item)

    def show_commodity(self):
        temp_list = sorted(list(self.commodities.keys()))
        for i in temp_list:
            print(i, self.commodities[i])
            # for j in self.commodities[i]:
            #     print(j)
    
    def show_demand(self):
        print(self.demand_info)

    # def show_commodity2(self):
    #     temp_list = sorted(list(self.commodities2.keys()))
    #     for i in temp_list:
    #         print(self.commodities2[i])
    #         # for j in self.commodities2[i]:
    #         #     print(j)

# project = Project("project.docx")
# project.show_info()
            