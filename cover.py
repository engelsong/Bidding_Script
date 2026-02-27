from re import A
from docx import Document
from docx.enum import text
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


class Cover:
    """
    通过project实例创建封面
    """
    
    parts = ['正本', '副本一', '副本二']
    sections_1 = ['资格证明文件', '商务技术文件']
    sections_2 = ['投标函部分', '技术标部分', '经济标部分', '商务标部分',]
    bidder = '投标人：中国海外经济合作有限公司'

    def __init__(self, project):
        self.project = project
        self.name = self.project.name
        self.code = self.project.code
        self.date = self.project.date
        self.doc = Document()

    def insert_para(self, word, size, font='黑体', space_after=Pt(0), alignment=text.WD_PARAGRAPH_ALIGNMENT.CENTER, 
    line_spacing=text.WD_LINE_SPACING.ONE_POINT_FIVE, left_indent=Pt(0)):
        """
        根据参数插入段落,word为文字,font是字体,size是字号,space_after是段后间距,alignment是对齐方式,line_spacing是行间距,
        left_indent是左缩进数据
        """

        para_now = self.doc.add_paragraph()
        para_now.paragraph_format.alignment = alignment
        para_now.paragraph_format.line_spacing_rule = line_spacing
        para_now.paragraph_format.left_indent = left_indent
        para_now.paragraph_format.space_after = space_after
        run_now = para_now.add_run(word)
        run_now.bold = True
        run_now.font.name = font
        run_now._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        run_now.font.size = Pt(size)
        return para_now

    def generate(self):
        """
        生成封面
        """
        for part in self.parts:
            for section in self.sections_1:                
                self.insert_para(part, 32, alignment=text.WD_PARAGRAPH_ALIGNMENT.RIGHT)
                for i in range(6):
                    self.insert_para('',  11)
                self.insert_para('投标文件', 48)
                self.insert_para('({})'.format(section), 32)
                for i in range(8):
                    self.insert_para('', 11)
                self.insert_para('项目名称:{}'.format(self.name), 18, left_indent=Pt(16), alignment=text.WD_PARAGRAPH_ALIGNMENT.LEFT)
                self.insert_para('招标编号：{}'.format(self.code), 18, left_indent=Pt(16), alignment=text.WD_PARAGRAPH_ALIGNMENT.LEFT)
                self.insert_para(self.bidder, 18, left_indent=Pt(16), alignment=text.WD_PARAGRAPH_ALIGNMENT.LEFT)
                self.insert_para('开标日期：{}'.format(self.date), 18, left_indent=Pt(16), alignment=text.WD_PARAGRAPH_ALIGNMENT.LEFT)
                self.doc.add_page_break()
        for section in self.sections_2:
            last = False
            if section == '商务标部分':
                last = True
            for i in range(12):
                self.insert_para('',  11)
            self.insert_para('商务技术文件', 48)
            self.insert_para('({})'.format(section), 32)
           
            if not last:
                self.doc.add_page_break()

        self.doc.save('封面-{}.docx'.format(self.name))

        
                



        
    



    



        